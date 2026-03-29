import asyncio
import shutil
import asyncpg
from docx import Document
import pandas as pd
from datetime import datetime
import re
import os
from openpyxl import load_workbook
from dateutil.relativedelta import relativedelta
from handlers.config import config
import math
import tempfile
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from openpyxl.utils import get_column_letter
from shutil import copy2
import time
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

async def get_data(query: str, *params):
    conn = None
    try:
        conn = await asyncpg.connect(config.db_connection)
        return await conn.fetch(query, *params)
    except Exception as e:
        print(f"Ошибка: {e}")
        return None
    finally:
        if conn:
            await conn.close()

async def get_info_business(id):
    us_businesses = await get_data('SELECT toa.name, fodb.name AS fodb_name, b.* FROM users u RIGHT JOIN Bussines b ON b.id = u.id_business JOIN Type_of_Activity toa ON b.id_type_of_activity = toa.id JOIN form_of_doing_business fodb ON fodb.id = b.id_form WHERE u.User_Id = $1', str(id))
    return us_businesses

async def delete_sheet_in_excel(sheet_name):
    try:
        file_path = 'docs/ГИРА_1006теккаа2.xlsx'
        loop = asyncio.get_event_loop()
        wb = await loop.run_in_executor(None, load_workbook, file_path)
        
        if sheet_name not in wb.sheetnames:
            raise ValueError(f"Лист '{sheet_name}' не найден")
        
        # Удаляем лист
        wb.remove(wb[sheet_name])
        
        # Сохраняем в отдельном потоке
        await loop.run_in_executor(None, wb.save, file_path)
        
        print(f"Лист '{sheet_name}' успешно удален")
        
    except FileNotFoundError:
        print(f"Файл {file_path} не найден")
    except Exception as e:
        print(f"Ошибка: {e}")

async def get_sheet_name_in_id_business(id_business):
    sheet_name = None
    results = await get_data('SELECT sheet_name FROM bussines WHERE id = $1',id_business)
    if results:
        for list_item in results:
            sheet_name = list_item['sheet_name']
    return sheet_name

async def delete_in_excel(name_tenant):
    wb = load_workbook('docs/ГИРА_1006теккаа2.xlsx')
    sheet = wb['Реестр']
    col_arendator = None
    for col in range(1, sheet.max_column + 1):
        if sheet.cell(row=1, column=col).value == "Арендатор":
            col_arendator = col
            break

    if col_arendator:
        search_value = name_tenant
        for row in range(sheet.max_row, 1, -1):
            if sheet.cell(row=row, column=col_arendator).value == search_value:
                sheet.delete_rows(row)
                print(f"Удалена строка {row} на листе 'Реестр'")
                break
        wb.save('docs/ГИРА_1006теккаа2.xlsx')
    else:
        print("Столбец 'Арендатор' не найден на листе 'Реестр'")

async def add_tenant_for_user(id, volume, amount, exploitation=0.0, unexpected=0.0):
    file_path_new = 'docs/ГИРА_1006теккаа2.xlsx'
    try:
        name_sheet = await get_sheet_name_in_id_business(id)
        df = pd.read_excel(file_path_new, sheet_name=name_sheet)
        wb = load_workbook(file_path_new)
        ws = wb[name_sheet]
        current_date = datetime.now()
        target_date = datetime(current_date.year, current_date.month, 1).date()
        
        col_pos = None
        for col in df.columns:
            try:
                if isinstance(col, datetime) and col.date() == target_date:
                    col_pos = df.columns.get_loc(col)
                    break
                elif isinstance(col, str) and pd.to_datetime(col).date() == target_date:
                    col_pos = df.columns.get_loc(col)
                    break
            except: continue
            
        if col_pos is not None:
            ws.cell(row=9, column=col_pos + 1, value=volume)
            ws.cell(row=9, column=col_pos + 3, value=amount)
            ws.cell(row=12, column=col_pos + 3, value=exploitation)
            ws.cell(row=13, column=col_pos + 3, value=unexpected)
            
            wb.save(file_path_new)
    except Exception as e:
        print(f'Ошибка в add_tenant_for_user: {e}')

async def safe_add_to_excel(data):
    file_path = 'docs/ГИРА_1006теккаа2.xlsx' 
    target_sheet = 'Реестр'
    if not os.path.exists(file_path): return
    try:
        wb = load_workbook(file_path)
        ws = wb[target_sheet]
        row = 2
        while ws[f'A{row}'].value is not None:
            row += 1
        ws[f'A{row}'] = data[0]
        ws[f'D{row}'] = data[1] 
        ws[f'E{row}'] = data[2]
        ws[f'F{row}'] = data[3]
        ws[f'G{row}'] = data[4]
        ws[f'H{row}'] = data[5]
        ws[f'I{row}'] = data[6]
        ws[f'K{row}'] = data[7]
        ws[f'R{row}'] = data[8]
        ws[f'S{row}'] = data[9]
        ws[f'T{row}'] = data[10]
        ws[f'V{row}'] = data[11]
        ws[f'N{row}'] = data[12]
        wb.save(file_path)
        wb.close()
    except Exception as e:
        print(f"Ошибка в safe_add_to_excel: {e}")

async def copy_sheet_safe(new_sheet_name):
    file_path = 'docs/ГИРА_1006теккаа2.xlsx'
    source_sheet_name = 'ЛИСТШАБЛОН'
    if not os.path.exists(file_path): return False
    try:
        wb = load_workbook(file_path)
        if source_sheet_name not in wb.sheetnames: return False
        if new_sheet_name in wb.sheetnames:
            base = new_sheet_name
            c = 1
            while new_sheet_name in wb.sheetnames:
                new_sheet_name = f"{base}_{c}"
                c += 1
        new_sheet = wb.copy_worksheet(wb[source_sheet_name])
        new_sheet.title = new_sheet_name
        wb.save(file_path)
        return True
    except Exception as e:
        print(f"Ошибка в copy_sheet_safe: {e}")
        return False

async def get_volume_and_amount_month(id_us):
    from handlers.meter_readings import get_sheet_name
    records = await get_info_business(id_us)
    square = records[0]['square'] if records else 0
    name_sheet = await get_sheet_name(id_us)
    file_path = 'docs/ГИРА_1006теккаа2.xlsx'
    df = pd.read_excel(file_path, sheet_name=name_sheet, engine='openpyxl')
    
    curr = datetime.now()
    ago = curr - relativedelta(months=1)
    dates = [datetime(curr.year, curr.month, 1).date(), datetime(ago.year, ago.month, 1).date()]
    cols = [df.columns.get_loc(c) for c in df.columns if isinstance(c, datetime) and c.date() in dates]
    if len(cols) < 2: return "Ошибка: данные не найдены."
    cols.sort()
    data = df.iloc[:, list(range(cols[0], cols[1] + 3))].values.tolist()
    
    res = []
    # Электро, ХВ, ГВ, Отопление, Экспл
    rows_idx = [2, 4, 6, 8, 10]
    for r_idx in rows_idx:
        row = data[r_idx]
        prev, curr_val = row[0] or 0, row[1] or 0
        amt = row[3] or 0
        rate = curr_val / amt if amt != 0 else 0
        res.append([prev, curr_val, curr_val, amt, rate])
    
    text = "📍 Счёт за прошлый месяц\n\n"
    # ... упрощенный вывод для целей примера, в реальности мы сохраняем логику
    return text

async def save_mr_result_in_excel(name_sheet, us_readings, type_id):
    file_path = 'docs/ГИРА_1006теккаа2.xlsx'
    wb = load_workbook(file_path)
    ws = wb[name_sheet]
    df = pd.read_excel(file_path, sheet_name=name_sheet)
    target = datetime(datetime.now().year, datetime.now().month, 1).date()
    col_pos = None
    for col in df.columns:
        if isinstance(col, datetime) and col.date() == target:
            col_pos = df.columns.get_loc(col)
            break
    if col_pos is not None:
        row = {1: 5, 2: 3, 3: 7}.get(type_id)
        try:
            val = int(us_readings)
        except (ValueError, TypeError):
            val = us_readings
        ws.cell(row=row, column=col_pos + 1, value=val)
    wb.save(file_path)

async def count_tenant_excel():
    file_path = 'docs/ГИРА_1006теккаа2.xlsx'
    # Читаем лист "Реестр"
    df = pd.read_excel(file_path, sheet_name='Реестр')

    # Считаем количество записей в столбце "Арендатор"
    count = df['Арендатор'].count()
    return count

async def create_excel(all_indicators,id_us,count_users):
    # ======== ПОЛУЧАЕМ ДАННЫЕ ИЗ ВАШЕГО КОДА ========
    from handlers.meter_readings import get_sheet_name
    
    # Получаем площадь
    records_list = await get_info_business(id_us)
    square = 0
    if records_list:
        for record in records_list:
            square = record['square']
    
    # Получаем имя листа
    name_sheet = await get_sheet_name(id_us)
    file_path = 'docs/ГИРА_1006теккаа2.xlsx'
    
    # ======== ЧИТАЕМ ДАННЫЕ ИЗ EXCEL ========
    # Используем pandas напрямую, xlwings на Linux не сработает
    df = pd.read_excel(file_path, sheet_name=name_sheet, engine='openpyxl')
    
    # Вычисляем даты
    current_date = datetime.now()
    date_month_ago = current_date - relativedelta(months=1)
    start_date = datetime(current_date.year, current_date.month, 1)
    end_date = datetime(date_month_ago.year, date_month_ago.month, 1)
    
    # Находим нужные колонки по датам
    found_columns = []
    for col in df.columns:
        col_date = None
        if isinstance(col, datetime):
            col_date = col
        if col_date and col_date.date() == start_date.date():
            found_columns.append(df.columns.get_loc(col))
        elif col_date and col_date.date() == end_date.date():
            found_columns.append(df.columns.get_loc(col))
    
    if len(found_columns) < 2:
        print("❌ Не найдены нужные колонки с датами для формирования отчета")
        return None
    start_index, end_index = found_columns[0], found_columns[1]
    col_range = list(range(start_index, end_index + 3))
    data_in_columns = df.iloc[:, col_range].values.tolist()
    
    # Формируем final_list с расчетами на стороне Python
    final_list = []
    
    # ⚡ ЭЛЕКТРОЭНЕРГИЯ (Строка 3 -> row_index 2)
    try:
        e_row = data_in_columns[2]
        e_prev, e_curr = e_row[0] or 0, e_row[1] or 0
        e_sum = e_row[3] or 0
        e_rate = e_curr / e_sum if e_sum != 0 else 0
        final_list.append([e_prev, e_curr, e_curr, e_sum, e_rate])
    except: final_list.append([0,0,0,0,0])

    # 🚰 ХВ (Строка 5, Ставка в Строке 6)
    try:
        cw_row = data_in_columns[4]
        cw_rate = data_in_columns[5][5] if len(data_in_columns[5]) > 5 else 0
        cw_prev, cw_curr = cw_row[0] or 0, cw_row[1] or 0
        cw_amt = cw_curr * cw_rate
        final_list.append([cw_prev, cw_curr, cw_curr, cw_amt, cw_rate])
    except: final_list.append([0,0,0,0,0])

    # 🔥 ГВ (Строка 7, Ставка в Строке 8)
    try:
        hw_row = data_in_columns[6]
        hw_rate = data_in_columns[7][5] if len(data_in_columns[7]) > 5 else 0
        hw_prev, hw_curr = hw_row[0] or 0, hw_row[1] or 0
        hw_amt = hw_curr * hw_rate
        final_list.append([hw_prev, hw_curr, hw_curr, hw_amt, hw_rate])
    except: final_list.append([0,0,0,0,0])

    # 🌡 ОТОПЛЕНИЕ (Берем из доп параметров или строки 9)
    try:
        heat_rate = data_in_columns[7][5] if len(data_in_columns[7]) > 5 else 0
        final_list.append([heat_rate * square, heat_rate])
    except: final_list.append([0, 0])

    # 🏢 ЭКСПЛУАТАЦИЯ (Строка 11, делим строго на 8)
    try:
        expl_total = data_in_columns[10][3] if len(data_in_columns[10]) > 3 else 0
        final_list.append([expl_total / 8])
    except: final_list.append([0])
    
    print(f"✅ Расчетный final_list готов: {final_list}")
    print(f'{all_indicators}')
    # ======== СОЗДАЕМ ВРЕМЕННЫЙ EXCEL ФАЙЛ ========
    temp_dir = tempfile.gettempdir()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    temp_file = os.path.join(temp_dir, f'report_{id_us}_{timestamp}.xlsx')
    
    # Создаем рабочую книгу
    wb = Workbook()
    
    # ======== СТИЛИ ========
    header_font = Font(bold=True, size=12)
    center_alignment = Alignment(horizontal='center', vertical='center')
    left_alignment = Alignment(horizontal='left', vertical='center')
    right_alignment = Alignment(horizontal='right', vertical='center')
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    header_fill = PatternFill(start_color='E6E6E6', end_color='E6E6E6', fill_type='solid')
    total_fill = PatternFill(start_color='FFFF99', end_color='FFFF99', fill_type='solid')
    
    # ======== ЛИСТ 1: ПОКАЗАТЕЛИ АДМИНА ========
    ws_admin = wb.active
    ws_admin.title = "Показатели ГИРА"
    
    # Заголовки для админа
    headers = ['Показатель', 'Объем', 'Сумма (руб.)']
    for col, header in enumerate(headers, 1):
        cell = ws_admin.cell(row=1, column=col, value=header)
        cell.font = header_font
        cell.alignment = center_alignment
        cell.border = thin_border
        cell.fill = header_fill
    
    # Данные админа
    admin_data = [
        ('⚡️ Электроэнергия', 
         all_indicators.get('electro', {}).get('volume', 0),
         all_indicators.get('electro', {}).get('amount', 0)),
        ('🏢 Экспл. услуги',
         '—',
         all_indicators.get('expl', {}).get('amount', 0)),
        ('💧 Водоотведение',
         '—',
         all_indicators.get('drainage', {}).get('amount', 0))
    ]

    count_users = int(count_users)
    expl_amount = float(all_indicators.get("expl", {}).get("amount", 0))
    price_for_user = expl_amount/count_users



    electro_amount = all_indicators['electro']['amount']
    water_cold_amount = all_indicators['water_cold']['amount']
    water_hot_amount = all_indicators['water_hot']['amount']
    drainage_amount = all_indicators['drainage']['amount']
    full_sum_amount = float(electro_amount)+float(water_cold_amount)+float(water_hot_amount)+float(expl_amount)+float(drainage_amount)
    for row, (name, volume, amount) in enumerate(admin_data, 2):
        cell = ws_admin.cell(row=row, column=1, value=name)
        cell.font = Font(bold=True)
        cell.alignment = left_alignment
        cell.border = thin_border
        
        cell = ws_admin.cell(row=row, column=2, value=volume)
        cell.alignment = center_alignment
        cell.border = thin_border
        if volume == '—':
            cell.value = '—'
        else:
            cell.number_format = '#,##0.00'
        
        cell = ws_admin.cell(row=row, column=3, value=float(amount) if amount != '—' else 0)
        cell.number_format = '#,##0.00'
        cell.alignment = center_alignment
        cell.border = thin_border
    
    # Итог для админа
    total_row = len(admin_data) + 2
    cell = ws_admin.cell(row=total_row, column=2, value="ИТОГО:")
    cell = ws_admin.cell(row=total_row, column=3, value=full_sum_amount)
    cell.font = Font(bold=True)
    cell.alignment = right_alignment
    cell.border = thin_border
    
    cell = ws_admin.cell(row=total_row, column=3, value=f"=SUM(C2:C{total_row-1})")
    cell.font = Font(bold=True)
    cell.number_format = '#,##0.00'
    cell.alignment = center_alignment
    cell.border = thin_border
    cell.fill = total_fill
    
    # ======== ЛИСТ 2: ПОКАЗАТЕЛИ ПОЛЬЗОВАТЕЛЯ ========
    ws_user = wb.create_sheet("Показатели пользователя")
    
    # Заголовки для пользователя
    user_headers = ['Услуга', 'Предыдущий', 'Текущий', 'Разница', 'Ставка', 'Сумма']
    for col, header in enumerate(user_headers, 1):
        cell = ws_user.cell(row=1, column=col, value=header)
        cell.font = header_font
        cell.alignment = center_alignment
        cell.border = thin_border
        cell.fill = header_fill
    sum_drainage = float(final_list[1][1]) + float(final_list[2][1])
    service_names = [
    ('⚡️ Электроэнергия', final_list[0] if len(final_list) > 0 else []),
    ('🚰 Холодная вода', final_list[1] if len(final_list) > 1 else []),
    ('🔥 Горячая вода', final_list[2] if len(final_list) > 2 else []),
    ('🌡 Отопление', final_list[3] if len(final_list) > 3 else []),
    ('🏢 Экспл. услуги', final_list[4] if len(final_list) > 4 else []),
    ('💧 Водоотведение', sum_drainage)
]

    row = 2
    formulas_row_start = row
    for service_name, data in service_names:
        if not data:
            continue
        
        cell = ws_user.cell(row=row, column=1, value=service_name)
        cell.font = Font(bold=True)
        cell.alignment = left_alignment
        cell.border = thin_border
        
        if service_name == '🌡 Отопление':
            if len(data) >= 2:
                sum = (data[0]) * float(data[1])
                ws_user.cell(row=row, column=2, value="—")
                ws_user.cell(row=row, column=3, value=data[0])
                ws_user.cell(row=row, column=4, value="—")
                ws_user.cell(row=row, column=5, value=data[1])
                ws_user.cell(row=row, column=6, value=f"=C{row}*E{row}")
                
        elif service_name == '🏢 Экспл. услуги' or service_name == '💧 Водоотведение':
            if service_name == '💧 Водоотведение':
                    # Для водоотведения - вставляем сумму в текущий показатель
                    drainage_amount = all_indicators['drainage']['amount']
                    sum_drainage = float(final_list[1][2]) + float(final_list[2][2])
                    final_drainage_amount = sum_drainage*float(drainage_amount)
                    ws_user.cell(row=row, column=3, value=sum_drainage)
                    ws_user.cell(row=row, column=5, value=drainage_amount)  # Текущий
                    ws_user.cell(row=row, column=6, value=final_drainage_amount)  # Сумма
            elif len(data) >= 1:
                # Для эксплуатации - как обычно
                ws_user.cell(row=row, column=6, value=price_for_user)  # Сумма
                ws_user.cell(row=row, column=2, value='—')  # Предыдущий
                ws_user.cell(row=row, column=3, value='—')  # Текущий
                ws_user.cell(row=row, column=4, value='—')  # Разница
                ws_user.cell(row=row, column=5, value='—')  # Ставка
        else:
            if len(data) >= 4:
                ws_user.cell(row=row, column=2, value=data[0])
                ws_user.cell(row=row, column=3, value=data[1])
                ws_user.cell(row=row, column=4, value=data[2])
                
                if len(data) >= 5:
                    ws_user.cell(row=row, column=5, value=data[4])
                    ws_user.cell(row=row, column=6, value=data[3])
                else:
                    ws_user.cell(row=row, column=5, value=data[2])
        
        # Форматирование
        for col in range(2, 7):
            cell = ws_user.cell(row=row, column=col)
            if cell.value != '—' and cell.value is not None and not str(cell.value).startswith('='):
                try:
                    float(cell.value)
                    cell.number_format = '#,##0.00'
                except:
                    pass
            cell.alignment = center_alignment
            cell.border = thin_border
        
        row += 1
    
    # Итог для пользователя
    total_row = row
    cell = ws_user.cell(row=total_row, column=5, value="ИТОГО К ОПЛАТЕ:")
    cell.font = Font(bold=True)
    cell.alignment = right_alignment
    cell.border = thin_border
    
    if row > formulas_row_start:
        cell = ws_user.cell(row=total_row, column=6, value=f"=SUM(F{formulas_row_start}:F{total_row-1})")
        cell.font = Font(bold=True)
        cell.number_format = '#,##0.00'
        cell.alignment = center_alignment
        cell.border = thin_border
        cell.fill = total_fill
    
    # Настройка ширины колонок
    column_widths = [30, 15, 15, 15, 15, 15, 15]
    for i, width in enumerate(column_widths, 1):
        ws_user.column_dimensions[get_column_letter(i)].width = width
    
    ws_admin.column_dimensions['A'].width = 30
    ws_admin.column_dimensions['B'].width = 15
    ws_admin.column_dimensions['C'].width = 15
    
    # Сохраняем файл
    wb.save(temp_file)
    wb.close()
    
    # На Linux мы не используем xlwings для пересчета в финальном файле.
    # Формулы будут пересчитаны на стороне пользователя при открытии в Excel.
    
    print(f"✅ Временный файл создан: {temp_file}")
    return temp_file


def cleanup_temp_file(file_path: str):
    """Удаляет временный файл"""
    try:
        if os.path.exists(file_path):
            os.remove(file_path)
            print(f"✅ Временный файл удален: {file_path}")
    except Exception as e:
        print(f"❌ Ошибка при удалении файла: {e}")


async def admin_indicators(all_indicators):
    file_path = 'docs/ГИРА_1006теккаа2.xlsx'
    name_sheet = '0. ГИРА'
    wb = load_workbook(file_path)
    ws = wb[name_sheet]
    df = pd.read_excel(file_path, sheet_name=name_sheet)
    ago = datetime.now() - relativedelta(months=1)
    target = datetime(ago.year, ago.month, 1).date()
    col_pos = None
    for col in df.columns:
        if isinstance(col, datetime) and col.date() == target:
            col_pos = df.columns.get_loc(col)
            break
    if col_pos is not None:
        # Электро: объём и сумма
        ws.cell(row=2, column=col_pos - 1, value=all_indicators['electro']['volume'])
        ws.cell(row=2, column=col_pos + 1, value=all_indicators['electro']['amount'])
        # Электро: ставка (руб/кВт·ч) — row 3 = pandas row 2 (header=None), откуда mast['electro'] читает
        electro_rate = all_indicators.get('electro', {}).get('tariff')
        if electro_rate is not None:
            ws.cell(row=3, column=col_pos + 1, value=electro_rate)
        # Вода ставка
        ws.cell(row=5, column=col_pos + 1, value=all_indicators['water_cold']['tariff'])
        ws.cell(row=7, column=col_pos + 1, value=all_indicators['water_cold']['tariff'])
        # Экспл
        ws.cell(row=10, column=col_pos + 1, value=all_indicators.get("expl", {}).get("amount", 0))
    wb.save(file_path)

async def compute_ku_total_from_excel(id_us, ind=None, unexp=0.0):
    """Compute the total KU amount for a tenant from the shared Excel file.

    Returns the total as a float, or None if the required date columns are missing.
    """
    from handlers.meter_readings import get_sheet_name
    if ind is None:
        ind = {}

    sheet = await get_sheet_name(id_us)
    path = os.path.join(os.getcwd(), 'docs', 'ГИРА_1006теккаа2.xlsx')

    mast = {'electro': 0.0, 'water_cold': 0.0, 'water_hot': 0.0}
    df_m = pd.read_excel(path, sheet_name='0. ГИРА', header=None).fillna(0)
    for i, k in {2: 'electro', 5: 'water_cold', 7: 'water_hot'}.items():
        if i < len(df_m):
            row = df_m.iloc[i].tolist()
            for v in reversed(row):
                try:
                    if float(v) > 0:
                        mast[k] = float(v)
                        break
                except:
                    continue

    df = pd.read_excel(path, sheet_name=sheet).fillna(0)
    curr, ago = datetime.now(), datetime.now() - relativedelta(months=1)
    trg = [datetime(curr.year, curr.month, 1).date(), datetime(ago.year, ago.month, 1).date()]
    cls = sorted([df.columns.get_loc(c) for c in df.columns if isinstance(c, datetime) and c.date() in trg])
    if len(cls) < 2:
        return None

    data = df.iloc[:, list(range(cls[0], cls[1] + 3))].values.tolist()
    final = []
    for i in range(2, 11, 2):
        r = data[i - 1]
        prev_v, curr_v = float(r[0]), float(r[3])
        rate_v = float(data[i][5])
        if rate_v == 0:
            key = {2: 'electro', 4: 'water_cold', 6: 'water_hot'}.get(i)
            rate_v = mast.get(key, 0)
        sum_v = float(r[5]) or round(curr_v * rate_v, 2)
        final.append([prev_v, curr_v, curr_v, sum_v, rate_v])

    expl_v = float(data[10][5]) if len(data) > 10 and len(data[10]) > 5 else round(float(ind.get('expl', {}).get('amount', 0)) / 8, 2)
    unexp_v = float(data[11][5]) if len(data) > 11 and len(data[11]) > 5 else float(unexp)
    dr_rate = float(ind.get('drainage', {}).get('amount', 0))
    dr_sum = round((final[1][2] + final[2][2]) * dr_rate, 2)

    sums = [final[0][3], final[1][3], final[2][3], final[3][3], expl_v, dr_sum, unexp_v]
    return round(sum(sums), 2)


async def create_word(*args, **kwargs):
    from handlers.meter_readings import get_sheet_name
    import traceback
    
    # Args: id_us, all_indicators, count_users, all_information, [unexpected]
    id_us, ind, cnt_u, info, unexp = None, {}, 0, [], 0.0
    if len(args) == 4: ind, id_us, cnt_u, info = args
    elif len(args) == 5: ind, id_us, cnt_u, info, unexp = args
    elif len(args) >= 7: id_us, cnt_u, ind, info, unexp = args[0], args[1], args[2], args[3], args[4]
    else: return (None, 0)

    try:
        rec = await get_info_business(id_us)
        full_n = f"{rec[0]['fodb_name']} {rec[0]['name_company']}" if rec else "—"
        short_n = f"{rec[0]['surname']} {rec[0]['first_name'][0]}.{rec[0]['patronymic'][0]}." if rec else "—"
        square = rec[0]['square'] if rec else 0
        id_form = rec[0]['id_form'] if rec else None
        
        sheet = await get_sheet_name(id_us)
        path = os.path.join(os.getcwd(), 'docs', 'ГИРА_1006теккаа2.xlsx')
        if id_form == 1:
            tpl = os.path.join(os.getcwd(), 'docs', 'Акт_расчета_ООО.docx')
        else:
            tpl = os.path.join(os.getcwd(), 'docs', 'Акт_расчета_ИП.docx')
        
        # Читаем ставки из 0. ГИРА
        mast = {'electro': 0.0, 'water_cold': 0.0, 'water_hot': 0.0}
        df_m = pd.read_excel(path, sheet_name='0. ГИРА', header=None).fillna(0)
        for i, k in {2: 'electro', 5: 'water_cold', 7: 'water_hot'}.items():
            if i < len(df_m):
                row = df_m.iloc[i].tolist()
                for v in reversed(row):
                    try: 
                        if float(v) > 0: mast[k] = float(v); break
                    except: continue

        df = pd.read_excel(path, sheet_name=sheet).fillna(0)
        curr, ago = datetime.now(), datetime.now() - relativedelta(months=1)
        trg = [datetime(curr.year, curr.month, 1).date(), datetime(ago.year, ago.month, 1).date()]
        cls = sorted([df.columns.get_loc(c) for c in df.columns if isinstance(c, datetime) and c.date() in trg])
        if len(cls) < 2: return (None, 0)
        
        data = df.iloc[:, list(range(cls[0], cls[1] + 3))].values.tolist()
        final = []
        for i in range(2, 11, 2): # Electro, CW, HW, Heat, Expl
            r = data[i-1] # Row in data
            prev_v, curr_v = float(r[0]), float(r[3])
            rate_v = float(data[i][5])
            if rate_v == 0:
                key = {2: 'electro', 4: 'water_cold', 6: 'water_hot'}.get(i)
                rate_v = mast.get(key, 0)
            sum_v = float(r[5]) or round(curr_v * rate_v, 2)
            final.append([prev_v, curr_v, curr_v, sum_v, rate_v])
        
        # Individual values
        expl_v = float(data[10][5]) if len(data) > 10 and len(data[10]) > 5 else round(float(ind.get('expl', {}).get('amount', 0))/8, 2)
        unexp_v = float(data[11][5]) if len(data) > 11 and len(data[11]) > 5 else float(unexp)
        dr_rate = float(ind.get('drainage', {}).get('amount', 0))
        dr_sum = round((final[1][2] + final[2][2]) * dr_rate, 2)

        fd, res_p = tempfile.mkstemp(suffix='.docx')
        os.close(fd)
        copy2(tpl, res_p)
        doc = Document(res_p)
        
        sub = {
            'start_day': str(info[1]), 'thisday': datetime.now().strftime("%d.%m.%Y"),
            'end_day': str(info[2]), 'monthyear': str(info[3]), 'square': str(square), 'tenantshortsfp': short_n
        }
        if id_form == 1:
            DIRECTOR_TITLE_GEN = {
                'Директор': 'директора',
                'Генеральный директор': 'генерального директора',
            }
            raw_title = rec[0].get('director_title') or 'Директор' if rec else 'Директор'
            sub['tenant_company'] = rec[0]['name_company'] if rec else '—'
            sub['tenant_director_title'] = DIRECTOR_TITLE_GEN.get(raw_title, raw_title.lower())
            sub['tenant_fio'] = f"{rec[0]['surname']} {rec[0]['first_name']} {rec[0]['patronymic']}" if rec else '—'
        else:
            sub['tenant_fio'] = f"{rec[0]['surname']} {rec[0]['first_name']} {rec[0]['patronymic']}" if rec else '—'
        
        def fmt(run, s=12, b=False):
            run.font.name = 'Times New Roman'
            run.font.size = Pt(s)
            run.bold = b
            r = run._element.get_or_add_rPr().get_or_add_rFonts()
            r.set(qn('w:ascii'), 'Times New Roman')
            r.set(qn('w:hAnsi'), 'Times New Roman')

        for p in list(doc.paragraphs) + [p for t in doc.tables for r in t.rows for c in r.cells for p in c.paragraphs]:
            for k, v in sub.items():
                if k in p.text: p.text = p.text.replace(k, v)
            for r in p.runs: fmt(r)

        tot = 0
        for p in doc.paragraphs:
            if 'Table_readings' in p.text:
                par = p._element.getparent()
                placeholder_idx = list(par).index(p._element)
                par.remove(p._element)
                tbl = doc.add_table(rows=9, cols=6)
                for c in tbl._element.xpath('.//w:tc'):
                    b = OxmlElement('w:tcBorders')
                    for s in ['top', 'left', 'bottom', 'right']:
                        e = OxmlElement(f'w:{s}')
                        e.set(qn('w:val'), 'single'), e.set(qn('w:sz'), '4'), e.set(qn('w:color'), '000000')
                        b.append(e)
                    c.get_or_add_tcPr().append(b)
                
                h = ["Услуга", "Предыдущий", "Текущий", "Разница", "Ставка", "Сумма (руб.)"]
                for j, txt in enumerate(h):
                    tbl.cell(0, j).text = txt
                    for r in tbl.cell(0, j).paragraphs[0].runs: fmt(r, 14, True)
                
                rows = [
                    ["Электроэнергия", str(final[0][0]), str(final[0][1]), str(final[0][2]), str(final[0][4]), str(final[0][3])],
                    ["Холодная вода", str(final[1][0]), str(final[1][1]), str(final[1][2]), str(final[1][4]), str(final[1][3])],
                    ["Горячая вода", str(final[2][0]), str(final[2][1]), str(final[2][2]), str(final[2][4]), str(final[2][3])],
                    ["Отопление", "—", "—", "—", "—", str(final[3][3])],
                    ["Эксплуатация", "—", "—", "—", "—", str(expl_v)],
                    ["Водоотведение", "—", str(final[1][2]+final[2][2]), "—", str(dr_rate), str(dr_sum)],
                    ["Непредвиденные", "—", str(unexp_v), "—", "—", str(unexp_v)]
                ]
                for i, row in enumerate(rows, 1):
                    for j, val in enumerate(row):
                        tbl.cell(i, j).text = val
                        for r in tbl.cell(i, j).paragraphs[0].runs: fmt(r)
                
                tot = round(sum([float(r[5]) for r in rows if r[5] != "—"]), 2)
                itogo_label = "ИТОГО" if id_form == 1 else "ИТОГО К ОПЛАТЕ"
                tbl.cell(8, 0).text, tbl.cell(8, 5).text = itogo_label, str(tot)
                for cell in [tbl.cell(8,0), tbl.cell(8,5)]:
                    for r in cell.paragraphs[0].runs: fmt(r, 12, True)
                par.remove(tbl._element)
                par.insert(placeholder_idx, tbl._element)
                break
        doc.save(res_p)
        return (res_p, tot)
    except: traceback.print_exc(); return (None, 0)