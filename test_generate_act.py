"""Тестовая генерация Акта расчёта КУ для ООО и ИП шаблонов (без БД)."""
import os
from shutil import copy2
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


def fmt(run, s=12, b=False):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(s)
    run.bold = b
    r = run._element.get_or_add_rPr().get_or_add_rFonts()
    r.set(qn('w:ascii'), 'Times New Roman')
    r.set(qn('w:hAnsi'), 'Times New Roman')


rows_data = [
    ["Электроэнергия\n(кВт·ч)", "12300", "12456", "156", "6.50", "1014.0"],
    ["Холодная вода (м³)",     "245",   "260",   "15",  "45.27", "679.05"],
    ["Горячая вода (м³)",      "120",   "128",   "8",   "214.46", "1715.68"],
    ["Водоотведение (м³)",     "—",     "23",    "—",   "20.0",   "460.0"],
    ["Отопление",              "—",     "—",     "—",   "—",      "1500.0"],
    ["Коммунальные услуги",    "—",     "—",     "—",   "—",      "108.7"],
    ["Непредвиденные", "—",     "500.0", "—",   "—",      "500.0"],
]


def generate_act(template_path, output_path, sub, itogo_label):
    copy2(template_path, output_path)
    doc = Document(output_path)

    for p in list(doc.paragraphs) + [p for t in doc.tables for r in t.rows for c in r.cells for p in c.paragraphs]:
        for k, v in sub.items():
            if k in p.text:
                p.text = p.text.replace(k, v)
        for r in p.runs:
            fmt(r)

    for p in doc.paragraphs:
        if 'Table_readings' in p.text:
            par = p._element.getparent()
            placeholder_idx = list(par).index(p._element)
            par.remove(p._element)

            tbl = doc.add_table(rows=9, cols=6)
            for c in tbl._element.xpath('.//w:tc'):
                b = OxmlElement('w:tcBorders')
                for side in ['top', 'left', 'bottom', 'right']:
                    e = OxmlElement(f'w:{side}')
                    e.set(qn('w:val'), 'single')
                    e.set(qn('w:sz'), '4')
                    e.set(qn('w:color'), '000000')
                    b.append(e)
                c.get_or_add_tcPr().append(b)

            headers = ["Услуга", "Предыдущий", "Текущий", "Разница", "Тариф", "Сумма (руб.)"]
            for j, txt in enumerate(headers):
                tbl.cell(0, j).text = txt
                for r in tbl.cell(0, j).paragraphs[0].runs:
                    fmt(r, 14, True)

            for i, row in enumerate(rows_data, 1):
                for j, val in enumerate(row):
                    tbl.cell(i, j).text = val
                    for r in tbl.cell(i, j).paragraphs[0].runs:
                        fmt(r)

            tot = round(sum(float(r[5]) for r in rows_data if r[5] != "—"), 2)
            tbl.cell(8, 0).text = itogo_label
            tbl.cell(8, 5).text = str(tot)
            for cell in [tbl.cell(8, 0), tbl.cell(8, 5)]:
                for r in cell.paragraphs[0].runs:
                    fmt(r, 12, True)

            par.remove(tbl._element)
            par.insert(placeholder_idx, tbl._element)
            break

    doc.save(output_path)
    print(f"Сгенерирован: {output_path}")


# --- ООО ---
ooo_sub = {
    'tenant_company': 'Тест-Компания',
    'tenant_director_title': 'директора',
    'tenant_fio': 'Иванов Иван Иванович',
    'start_day': '01.03.2026',
    'thisday': '28.03.2026',
    'end_day': '31.03.2026',
    'monthyear': 'Март 2026',
    'square': '55.0',
    'tenantshortsfp': 'Иванов И.И.',
}

generate_act(
    template_path=os.path.join(os.getcwd(), 'docs', 'Акт_расчета_ООО.docx'),
    output_path=os.path.join(os.getcwd(), 'test_act_output_ooo.docx'),
    sub=ooo_sub,
    itogo_label='ИТОГО',
)

# --- ИП ---
ip_sub = {
    'tenant_fio': 'Петров Петр Петрович',
    'start_day': '01.03.2026',
    'thisday': '28.03.2026',
    'end_day': '31.03.2026',
    'monthyear': 'Март 2026',
    'square': '42.5',
    'tenantshortsfp': 'Петров П.П.',
}

generate_act(
    template_path=os.path.join(os.getcwd(), 'docs', 'Акт_расчета_ИП.docx'),
    output_path=os.path.join(os.getcwd(), 'test_act_output_ip.docx'),
    sub=ip_sub,
    itogo_label='ИТОГО К ОПЛАТЕ',
)
